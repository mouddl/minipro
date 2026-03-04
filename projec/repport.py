from docx import Document
from docx.shared import Inches
import pandas as pd
import os

# Create a new Document
doc = Document()
doc.add_heading('Mini-Projet: Planification Robuste sur Grille', level=1)

# Define the base path to the results directory
base_path = os.path.abspath(os.path.join(os.path.dirname(__file__), 'results'))

# Chapter 1: Grids and Algorithms
doc.add_heading('Chapitre 1: Présentation des Grilles et des Chemins', level=1)

# Section: Grids
doc.add_heading('1.1 Présentation des Grilles', level=2)
doc.add_paragraph(
    "Nous avons utilisé trois grilles de complexité différente pour évaluer les performances des algorithmes. "
    "Les grilles sont présentées ci-dessous, avec le point de départ (START) et le but (GOAL)."
)

# Add all three grids in one page
for grid_type in ['easy', 'medium', 'hard']:
    image_path = os.path.join(base_path, f'grid_{grid_type}_plain.png')
    try:
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(3.0))
        else:
            doc.add_paragraph(f"Image not found: {image_path}")
    except Exception as e:
        doc.add_paragraph(f"Error loading image {image_path}: {e}")

# Section: Paths Found by Algorithms
doc.add_heading('1.2 Chemins Trouvés par les Algorithmes', level=2)
doc.add_paragraph(
    "Les algorithmes UCS, Greedy, et A* ont trouvé les chemins suivants sur chaque grille. "
    "Les chemins sont superposés pour une comparaison visuelle."
)

for grid_type in ['easy', 'medium', 'hard']:
    image_path = os.path.join(base_path, f'grid_{grid_type}_all_algorithms.png')
    try:
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(4.5))
        else:
            doc.add_paragraph(f"Image not found: {image_path}")
    except Exception as e:
        doc.add_paragraph(f"Error loading image {image_path}: {e}")

# Section: Execution Time Comparison
doc.add_heading('1.3 Comparaison des Temps d\'Exécution', level=2)
doc.add_paragraph(
    "Le graphique ci-dessous montre les temps d'exécution des algorithmes sur chaque grille. "
    "A* est généralement plus rapide que UCS et Greedy, surtout sur les grilles complexes."
)

image_path = os.path.join(base_path, 'bar_execution_time.png')
try:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.0))
    else:
        doc.add_paragraph(f"Image not found: {image_path}")
except Exception as e:
    doc.add_paragraph(f"Error loading image {image_path}: {e}")

# Section: Explored Nodes Analysis
doc.add_heading('1.4 Analyse des Nœuds Explorés', level=2)
doc.add_paragraph(
    "Le graphique suivant montre le nombre de nœuds développés par chaque algorithme. "
    "A* explore généralement moins de nœuds que UCS et Greedy, ce qui le rend plus efficace."
)

image_path = os.path.join(base_path, 'bar_nodes_explored.png')
try:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.0))
    else:
        doc.add_paragraph(f"Image not found: {image_path}")
except Exception as e:
    doc.add_paragraph(f"Error loading image {image_path}: {e}")

# Chapter 2: Markov Analysis
doc.add_heading('Chapitre 2: Analyse Markovienne', level=1)

# Section: Impact of Epsilon
doc.add_heading('2.1 Impact de ε sur la Probabilité d\'Atteindre le But', level=2)
doc.add_paragraph(
    "Les graphiques suivants montrent l'impact du paramètre d'incertitude ε sur la probabilité d'atteindre le but. "
    "L'augmentation de ε réduit la probabilité d'atteindre le but."
)

image_path = os.path.join(base_path, 'graph_prob_vs_epsilon.png')
try:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.0))
    else:
        doc.add_paragraph(f"Image not found: {image_path}")
except Exception as e:
    doc.add_paragraph(f"Error loading image {image_path}: {e}")

# Section: Expected Time to Absorption
doc.add_heading('2.2 Temps Moyen Avant Absorption', level=2)
doc.add_paragraph(
    "Ce graphique montre le temps moyen avant absorption en fonction de ε. "
    "L'augmentation de ε augmente le temps moyen avant absorption."
)

image_path = os.path.join(base_path, 'graph_expected_time_vs_epsilon.png')
try:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.0))
    else:
        doc.add_paragraph(f"Image not found: {image_path}")
except Exception as e:
    doc.add_paragraph(f"Error loading image {image_path}: {e}")

# Section: Distribution Analysis
doc.add_heading('2.3 Évolution de π(n) pour ε = 0.1', level=2)
doc.add_paragraph(
    "Ce graphique montre l'évolution de la distribution π(n) pour ε = 0.1 sur la grille medium. "
    "Il illustre la probabilité d'être dans l'état GOAL après n étapes."
)

image_path = os.path.join(base_path, 'distribution_pi_medium_eps01.png')
try:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.0))
    else:
        doc.add_paragraph(f"Image not found: {image_path}")
except Exception as e:
    doc.add_paragraph(f"Error loading image {image_path}: {e}")

# Analysis
doc.add_heading('Analyse', level=1)
analysis_text = (
    "Les résultats montrent que:\n\n"
    "- A* est l'algorithme le plus efficace en termes de nœuds explorés et de temps d'exécution.\n"
    "- L'augmentation de ε réduit la probabilité d'atteindre le but et augmente le temps moyen avant absorption.\n"
    "- La modélisation Markovienne permet de quantifier l'impact de l'incertitude sur la planification de chemin.\n\n"
    "Ces résultats soulignent l'importance de modéliser l'incertitude dans la planification de chemin."
)
doc.add_paragraph(analysis_text)

# Conclusion
doc.add_heading('Conclusion', level=1)
conclusion_text = (
    "En conclusion, ce projet a démontré l'efficacité de l'algorithme A* pour la planification de chemin "
    "et l'importance de modéliser l'incertitude avec les chaînes de Markov. "
    "Les résultats montrent un compromis clair entre la robustesse du plan et l'incertitude dans l'environnement."
)
doc.add_paragraph(conclusion_text)

# Save the document
doc_path = os.path.abspath('MiniProjet_PlanificationRobuste.docx')
doc.save(doc_path)

doc_path
